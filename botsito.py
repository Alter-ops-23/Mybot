import os
import telebot
from telebot.types import InputFile
from docx import Document
from docx.shared import Inches
import requests
import tempfile
import os
from datetime import datetime

# ---------------------------------------------------
# CONFIGURA TU TOKEN AQUI
# ---------------------------------------------------
BOT_TOKEN = os.getenv("BOT_TOKEN")
bot = telebot.TeleBot(BOT_TOKEN)

# ---------------------------------------------------
# MEMORIA TEMPORAL DE USUARIOS
# ---------------------------------------------------
usuarios = {}

# ---------------------------------------------------
# LISTA DE INTEGRANTES BASE
# ---------------------------------------------------
INTEGRANTES_ORIGINALES = [
    "ARREDONDO HERNÁNDEZ JONATHAN AXEL",
    "GALINDO CETINA LEONARDO AZAEL",
    "PEREZ BARRIENTOS AXEL NICOLAS BRYAN",
    "ROMANO GARCÍA JOSÉ EDUARDO",
    "VILLARRUEL SÁNCHEZ SOFÍA"
]

# ---------------------------------------------------
# CARGA LA PLANTILLA
# ---------------------------------------------------
PLANTILLA_PATH = "PLANTILLA_BLITZ.docx"

# ---------------------------------------------------
# INICIO DEL FLUJO
# ---------------------------------------------------
@bot.message_handler(commands=['start'])
def start(msg):
    bot.reply_to(msg, "HOLA, ESCRIBE: GENERAR REPORTE")

@bot.message_handler(func=lambda m: m.text and m.text.lower() == "generar reporte")
def generar_reporte(msg):
    uid = msg.from_user.id
    usuarios[uid] = {
        "asistencia": None,
        "faltantes": None,
        "fecha": datetime.now().strftime("%d/%m/%Y").upper(),
        "actividad": None,
        "evidencia": None,
        "archivo": None
    }
    bot.reply_to(msg, "¿ASISTIERON TODOS LOS INTEGRANTES? (SI / NO)")

# ---------------------------------------------------
# RESPUESTA SI / NO
# ---------------------------------------------------
@bot.message_handler(func=lambda m: m.from_user.id in usuarios and m.text.lower() in ["si", "no"])
def asistio(msg):
    uid = msg.from_user.id
    usuarios[uid]["asistencia"] = msg.text.upper()

    if msg.text.lower() == "si":
        bot.reply_to(msg, "PERFECTO. ¿QUÉ NOMBRE LLEVARÁ LA ACTIVIDAD?")
    else:
        bot.reply_to(msg, "¿QUIÉN O QUIÉNES FALTARON? (ESCRIBE LOS NOMBRES EXACTOS)")

# ---------------------------------------------------
# RECIBIR FALTANTES
# ---------------------------------------------------
@bot.message_handler(func=lambda m: m.from_user.id in usuarios and usuarios[m.from_user.id]["asistencia"] == "NO" and usuarios[m.from_user.id]["faltantes"] is None)
def faltantes(msg):
    uid = msg.from_user.id
    usuarios[uid]["faltantes"] = msg.text.upper()
    bot.reply_to(msg, "INTEGRANTES ACTUALIZADOS.\n¿QUÉ NOMBRE LLEVARÁ LA ACTIVIDAD?")

# ---------------------------------------------------
# RECIBIR NOMBRE DE ACTIVIDAD
# ---------------------------------------------------
@bot.message_handler(func=lambda m: m.from_user.id in usuarios and usuarios[m.from_user.id]["actividad"] is None and m.text.lower() not in ["si", "no"])
def actividad(msg):
    uid = msg.from_user.id
    usuarios[uid]["actividad"] = msg.text.upper()
    bot.reply_to(msg, "ENVÍA LA EVIDENCIA (FOTO O IMAGEN).")

# ---------------------------------------------------
# RECIBIR FOTO
# ---------------------------------------------------
@bot.message_handler(content_types=['photo'])
def evidencia(msg):
    uid = msg.from_user.id
    if uid not in usuarios:
        bot.reply_to(msg, "PRIMERO ESCRIBE: GENERAR REPORTE")
        return

    file_id = msg.photo[-1].file_id
    usuarios[uid]["evidencia"] = file_id
    bot.reply_to(msg, "¿QUÉ NOMBRE TENDRÁ EL ARCHIVO? (DEBE TERMINAR EN _BLITZ)")

# ---------------------------------------------------
# RECIBIR NOMBRE DE ARCHIVO
# ---------------------------------------------------
@bot.message_handler(func=lambda m: m.from_user.id in usuarios and usuarios[m.from_user.id]["archivo"] is None)
def nombre_archivo(msg):
    uid = msg.from_user.id
    nombre = msg.text.upper()

    if not nombre.endswith("_BLITZ"):
        bot.reply_to(msg, "EL NOMBRE DEBE TERMINAR EN _BLITZ. INTÉNTALO DE NUEVO.")
        return

    usuarios[uid]["archivo"] = nombre
    bot.reply_to(msg, "GENERANDO DOCUMENTO, ESPERA...")

    generar_documento(uid, msg)

# ---------------------------------------------------
# FUNCIÓN PRINCIPAL PARA CREAR EL DOCUMENTO
# ---------------------------------------------------
def generar_documento(uid, msg):
    data = usuarios[uid]

    # --- CARGAR DOCUMENTO ---
    doc = Document(PLANTILLA_PATH)

    # ---- PROCESAR INTEGRANTES ----
    if data["asistencia"] == "SI":
        lista_final = INTEGRANTES_ORIGINALES
    else:
        faltantes = [f.strip() for f in data["faltantes"].split(",")]
        lista_final = [
            i for i in INTEGRANTES_ORIGINALES
            if all(f not in i for f in faltantes)
        ]

    # ---- REMPLAZOS EN EL DOCUMENTO ----
    for p in doc.paragraphs:
        if "NOMBRE DE LOS INTEGRANTES" in p.text:
            idx = doc.paragraphs.index(p)
            for i in range(len(INTEGRANTES_ORIGINALES)):
                doc.paragraphs[idx + 2 + i].text = ""

            for n, nombre in enumerate(lista_final):
                doc.paragraphs[idx + 2 + n].text = nombre

        if "FECHA DE ENTREGA" in p.text:
            idx = doc.paragraphs.index(p)
            doc.paragraphs[idx + 2].text = data["fecha"]

        if "NOMBRE DE LA ACTIVIDAD" in p.text:
            idx = doc.paragraphs.index(p)
            doc.paragraphs[idx + 2].text = data["actividad"]

    # ---- DESCARGAR EVIDENCIA ----
    file_info = bot.get_file(data["evidencia"])
    url = f"https://api.telegram.org/file/bot{BOT_TOKEN}/{file_info.file_path}"
    img_data = requests.get(url).content

    tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
    tmp_img.write(img_data)
    tmp_img.close()

    # ---- NUEVA PÁGINA + FOTO ----
    doc.add_page_break()
    doc.add_picture(tmp_img.name, width=Inches(5.5))

    # ---- GUARDAR DOCUMENTO ----
    output = f"{data['archivo']}.docx"
    doc.save(output)

    # ---- ENVIAR ----
    bot.send_document(msg.chat.id, InputFile(output), caption="AQUÍ ESTÁ TU ARCHIVO.")

    # ---- LIMPIAR ----
    os.remove(output)
    os.remove(tmp_img.name)
    del usuarios[uid]

# ---------------------------------------------------
# INICIAR BOT
# ---------------------------------------------------
bot.polling()

